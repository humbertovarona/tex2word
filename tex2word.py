import sys
import matplotlib
matplotlib.use('Agg')

import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches

def generate_word_document(equation_latex):
    doc = Document()
    doc.add_heading('Equation in LaTeX', level=1)
    doc.add_paragraph("LaTeX Equation:")
    doc.add_paragraph(equation_latex)
    try:
        fig, ax = plt.subplots()
        ax.text(0.5, 0.5, f"${equation_latex}$", fontsize=20, ha='center', va='center')
        ax.axis('off')
        image_stream = BytesIO()
        plt.savefig(image_stream, format='png', bbox_inches='tight')
        image_stream.seek(0)  # Rewind the buffer
        doc.add_paragraph("Equation in Word (Image Format):")
        doc.add_picture(image_stream, width=Inches(4))
        plt.close(fig)
    except Exception as e:
        doc.add_paragraph(f"Error rendering LaTeX: {str(e)}")
    doc_name = "equation_output.docx"
    doc.save(doc_name)
    print(f"Document '{doc_name}' has been created with the equation.")


if len(sys.argv) < 2:
    print("Please provide a LaTeX equation.")
    sys.exit(1)
latex_equation = sys.argv[1]
latex_equation = f"{{{{{latex_equation}}}}}"
generate_word_document(latex_equation)
